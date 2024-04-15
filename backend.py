import pdfplumber
import docx
#import spire.doc
import re
import xlsxwriter
#from io import BytesIO


def extract_text_from_pdf(pdf_path):
    pdf = pdfplumber.open(pdf_path)
    pages = []
    for page in pdf.pages:
        page_text = page.extract_text()
        pages.append(page_text)
    pdf.close()
    text = "\n\n".join(pages)
    text = text.replace("\n"," ")
    return text


def extract_text_from_docx(doc_path):
    doc = docx.Document(doc_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    text = text.replace("\n", " ")
    return text


def extract_text_from_doc(doc_path):
    doc = spire.doc.Document()
    doc.LoadFromFile(doc_path)
    #text = doc.GetText()
    text = ''
    for paragraph in doc.paragraph:
        text += paragraph.text + '\n'
    text = text.replace("\n", " ")
    return text


def extract_email_ids(text):
    email_ids = ""
    for match in re.finditer(r"\S+@\S+", text):
        email_ids += match.group()
    return email_ids


def extract_phone_numbers(text):
    phone_numbers = ""
    for match in re.finditer(r"\d{10}", text):
        phone_numbers += match.group()
    return phone_numbers


def save_to_excel(data, output_file):
    workbook = xlsxwriter.Workbook(output_file)
    worksheet = workbook.add_worksheet()
    worksheet.write(0, 0, 'Email')
    worksheet.write(0, 1, 'Phone Number')
    worksheet.write(0, 2, 'Text')

    for row, (email, phone, text) in enumerate(data):
        worksheet.write(row + 1, 0, email)
        worksheet.write(row + 1, 1, phone)
        worksheet.write(row + 1, 2, text)

    workbook.close()
    print(f"Data saved to {output_file}")